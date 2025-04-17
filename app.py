from flask import Flask, render_template, request, send_file
from openpyxl import load_workbook
from docx import Document
from datetime import datetime
import io
import os
import pandas as pd

app = Flask(__name__)

CAMINHO_BASE = os.path.dirname(__file__)
ARQUIVO_EXCEL = os.path.join(CAMINHO_BASE, 'cooperados.xlsx')
MODELO_PF_AGRO = os.path.join(CAMINHO_BASE, 'modelo.docx')
MODELO_PJ = os.path.join(CAMINHO_BASE, 'modelo_pj.docx')

def carregar_dados(nome):
    wb = load_workbook(ARQUIVO_EXCEL, data_only=True)
    ws = wb.active
    dados = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        dados.append({
            'Nome': row[0],
            'Estado Civil': row[1],
            'Ocupação': row[2],
            'CPF/CNPJ': row[3],
            'Endereço': row[4],
            'CEP': row[5],
            'Cidade': row[6]
        })
    return pd.DataFrame(dados)

def substituir_campos(doc, substituicoes):
    for p in doc.paragraphs:
        for chave, valor in substituicoes.items():
            if chave in p.text:
                p.text = p.text.replace(chave, valor)

    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for chave, valor in substituicoes.items():
                    if chave in celula.text:
                        celula.text = celula.text.replace(chave, valor)

def formatar_cpf(cpf):
    cpf = ''.join(filter(str.isdigit, str(cpf)))
    return f'{cpf[:3]}.{cpf[3:6]}.{cpf[6:9]}-{cpf[9:]}' if len(cpf) == 11 else cpf

def formatar_cnpj(cnpj):
    cnpj = ''.join(filter(str.isdigit, str(cnpj)))
    return f'{cnpj[:2]}.{cnpj[2:5]}.{cnpj[5:8]}/{cnpj[8:12]}-{cnpj[12:]}' if len(cnpj) == 14 else cnpj

def formatar_cep(cep):
    cep = ''.join(filter(str.isdigit, str(cep)))
    return f'{cep[:5]}-{cep[5:]}' if len(cep) == 8 else cep

def formatar_rg(rg):
    rg = ''.join(filter(str.isdigit, str(rg)))
    return f'{rg[:2]}.{rg[2:5]}.{rg[5:]}' if len(rg) >= 7 else rg

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        tipo = request.form['tipo']
        df = carregar_dados(request.form['nome'])

        if tipo in ['PF', 'AGRO']:
            nome = request.form['nome'].strip()
            linha = df[df['Nome'].str.strip().str.lower() == nome.lower()]
            if linha.empty:
                return "Cooperado não encontrado na base de dados.", 404

            dados = linha.iloc[0]
            campos_pf = ['RG', 'Apelido do Dispositivo', 'Modelo do Dispositivo',
                         'Local', 'Nome do Colaborador', 'CPF do Colaborador']
            respostas = {campo: request.form.get(campo, '') for campo in campos_pf}

            agora = datetime.now()

            substituicoes = {
                'NOMECOOPERADO': dados['Nome'].strip(),
                'ESTADOCIVIL': dados['Estado Civil'].strip(),
                'OCUPACAO': dados['Ocupação'].strip(),
                'CPFCOOPERADO': formatar_cpf(dados['CPF/CNPJ']),
                'ENDERECO': dados['Endereço'].strip(),
                'CEP': formatar_cep(dados.get('CEP', '')),
                'CIDADE': dados['Cidade'].strip(),
                'DATA': agora.strftime('%d/%m/%Y'),
                'HORA': agora.strftime('%H:%M'),
                'RGCOOPERADO': formatar_rg(respostas['RG']),
                'APELIDODISPOSITIVO': respostas['Apelido do Dispositivo'],
                'MODELODISPOSITIVO': respostas['Modelo do Dispositivo'],
                'LOCAL': respostas['Local'],
                'NOMECOLABORADOR': respostas['Nome do Colaborador'],
                'CPFCOLABORADOR': formatar_cpf(respostas['CPF do Colaborador']),
            }

            doc = Document(MODELO_PF_AGRO)
            substituir_campos(doc, substituicoes)

        elif tipo == 'PJ':
            nome = request.form['nome'].strip()
            linha = df[df['Nome'].str.strip().str.lower() == nome.lower()]
            if linha.empty:
                return "Cooperado não encontrado na base de dados.", 404

            dados = linha.iloc[0]
            campos_pj = ['RG', 'CNPJ', 'Empresa', 'Chave Multicanal', 'Local', 'Nome do Colaborador', 'CPF do Colaborador']
            respostas = {campo: request.form.get(campo, '') for campo in campos_pj}

            agora = datetime.now()

            substituicoes = {
                'NOMECOOPERADO': dados['Nome'].strip(),
                'ESTADOCIVIL': dados['Estado Civil'].strip(),
                'OCUPACAO': dados['Ocupação'].strip(),
                'PESSOAJURIDICA': formatar_cnpj(respostas['CNPJ']),
                'EMPRESA': respostas['Empresa'],
                'CHAVE': respostas['Chave Multicanal'],
                'ENDERECO': dados['Endereço'].strip(),
                'CEP': formatar_cep(dados.get('CEP', '')),
                'CIDADE': dados['Cidade'].strip(),
                'DATA': agora.strftime('%d/%m/%Y'),
                'HORA': agora.strftime('%H:%M'),
                'RGCOOPERADO': formatar_rg(respostas['RG']),
                'LOCAL': respostas['Local'],
                'NOMECOLABORADOR': respostas['Nome do Colaborador'],
                'CPFCOLABORADOR': formatar_cpf(respostas['CPF do Colaborador']),
            }

            doc = Document(MODELO_PJ)
            substituir_campos(doc, substituicoes)

        output = io.BytesIO()
        doc.save(output)
        output.seek(0)

        return send_file(output, as_attachment=True, download_name="termo.docx")

    return render_template('index.html')

if __name__ == '__main__':
    app.run(debug=True)
